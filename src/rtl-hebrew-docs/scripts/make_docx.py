#!/usr/bin/env python3
"""Generate an EDITABLE Hebrew Word (.docx) file with correct RTL.

Choose this path ONLY when the recipient must edit the document in Microsoft
Word. For RTL to survive you must set bidi at three levels: the section, every
paragraph, and every run (with w:rtl AND w:cs for the complex-script font slot).

CAVEAT (learned the hard way): python-docx RTL is reliably honored only by
Microsoft Word. Pages, macOS Quick Look, Google Drive preview, and LibreOffice
frequently ignore the markers and show the text left-to-right. If the recipient
does not specifically need an editable Word file, deliver a PDF instead
(see html_to_pdf.py), it renders identically everywhere.

Usage:
    python3 make_docx.py --type contract --output contract.docx
    python3 make_docx.py --type letter   --output letter.docx

Requirements:
    pip install python-docx
"""

import argparse
import sys

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("Missing dependency. Install with: pip install python-docx")

HEBREW_FONT = "David"   # has Hebrew glyphs; "Heebo"/"Arial" also work. NOT Roboto.


def set_rtl_paragraph(paragraph):
    """Mark a paragraph as RTL and right-aligned."""
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(OxmlElement("w:bidi"))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def set_rtl_run(run):
    """Mark a run as RTL, including the complex-script slot (w:cs)."""
    rPr = run._r.get_or_add_rPr()
    rPr.append(OxmlElement("w:rtl"))
    rPr.append(OxmlElement("w:cs"))


def set_complex_script_font(run, font_name, size_pt):
    """Set the font on the complex-script slot so Hebrew uses the right face."""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:cs"), font_name)      # complex script (Hebrew) font
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def add_hebrew_paragraph(doc, text, size=12, bold=False, heading=False):
    """Add a fully RTL Hebrew paragraph and return it."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold or heading
    set_complex_script_font(run, HEBREW_FONT, 15 if heading else size)
    set_rtl_run(run)
    set_rtl_paragraph(para)
    return para


def set_section_rtl(doc):
    """Set the whole section's base direction to RTL."""
    sectPr = doc.sections[0]._sectPr
    bidi = sectPr.find(qn("w:bidi"))
    if bidi is None:
        sectPr.append(OxmlElement("w:bidi"))


def build_contract(doc):
    """A fictional service-agreement skeleton (safe to generate freely)."""
    add_hebrew_paragraph(doc, "הסכם התקשרות למתן שירותים", heading=True)
    add_hebrew_paragraph(
        doc,
        "שנערך ונחתם ביום ___ בחודש ___ שנת 2026, בין הצדדים כדלקמן:",
    )
    add_hebrew_paragraph(doc, "1. הצדדים", bold=True)
    add_hebrew_paragraph(doc, "המזמין: __________ (להלן: \"המזמין\").")
    add_hebrew_paragraph(doc, "נותן השירות: __________ (להלן: \"הספק\").")
    add_hebrew_paragraph(doc, "2. היקף השירותים", bold=True)
    add_hebrew_paragraph(doc, "הספק יספק למזמין את השירותים המפורטים בנספח א׳ להסכם זה.")
    add_hebrew_paragraph(doc, "3. תמורה ותנאי תשלום", bold=True)
    add_hebrew_paragraph(doc, "התמורה תשולם בתנאי שוטף + 30 כנגד חשבונית כדין.")
    add_hebrew_paragraph(doc, "4. תקופת ההסכם וסיומו", bold=True)
    add_hebrew_paragraph(doc, "כל צד רשאי לסיים הסכם זה בהודעה מוקדמת של 30 יום מראש ובכתב.")
    add_hebrew_paragraph(doc, "5. סודיות", bold=True)
    add_hebrew_paragraph(doc, "הצדדים מתחייבים לשמור על סודיות המידע שיימסר במסגרת ההתקשרות.")
    add_hebrew_paragraph(doc, "6. חתימות", bold=True)
    add_hebrew_paragraph(doc, "_______________               _______________")
    add_hebrew_paragraph(doc, "המזמין                                  הספק")


def build_letter(doc):
    """A fictional formal letter (safe to generate freely)."""
    add_hebrew_paragraph(doc, "סטודיו רקפת", heading=True)
    add_hebrew_paragraph(doc, "30 ביוני 2026")
    add_hebrew_paragraph(doc, "לכבוד: מחלקת הרכש")
    add_hebrew_paragraph(doc, "הנדון: אישור קבלת הזמנה", bold=True)
    add_hebrew_paragraph(
        doc,
        "שלום רב,\nברצוננו לאשר את קבלת הזמנתכם מיום 25 ביוני 2026. "
        "ההזמנה נקלטה במערכת ותטופל בהתאם ללוח הזמנים שסוכם.",
    )
    add_hebrew_paragraph(doc, "נשמח לעמוד לרשותכם בכל שאלה.")
    add_hebrew_paragraph(doc, "בברכה,\nצוות סטודיו רקפת")


BUILDERS = {"contract": build_contract, "letter": build_letter}


def main():
    parser = argparse.ArgumentParser(
        description="Generate an editable Hebrew (RTL) Word document."
    )
    parser.add_argument("--type", choices=sorted(BUILDERS), default="contract",
                        help="Document type to build (default: contract).")
    parser.add_argument("--output", default="document.docx", help="Output .docx path.")
    args = parser.parse_args()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = HEBREW_FONT
    style.font.size = Pt(12)
    set_section_rtl(doc)

    BUILDERS[args.type](doc)
    doc.save(args.output)

    print(f"Generated: {args.output}")
    print("VERIFY: open this in Microsoft Word, its RTL is NOT reliably shown "
          "by Pages / Quick Look / Google Drive preview / LibreOffice. "
          "If the recipient does not need to edit in Word, prefer a PDF.")


if __name__ == "__main__":
    main()
